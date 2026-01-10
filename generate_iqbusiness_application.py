#!/usr/bin/env python3
"""
Generate tailored CV and Cover Letter for IQbusiness React.js Software Developer position
Focuses on React.js, Node.js, TypeScript, JavaScript, and SQL requirements
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import datetime

def add_hyperlink(paragraph, url, text):
    """Add hyperlink to paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0563C1")
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_iqbusiness_cv():
    """Generate tailored CV for IQbusiness React.js Developer position"""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("KUTLOANO PETER MOSHAO")
    run.font.size = Pt(16)
    run.bold = True
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("Maseru, Lesotho | kutloano.moshao111@gmail.com | +266 5758 6176\\n")
    add_hyperlink(contact, "https://kutloano-showcase.vercel.app", "kutloano-showcase.vercel.app")
    contact.add_run(" | ")
    add_hyperlink(contact, "https://linkedin.com/in/kutloano-moshao-1aa5003a1", "LinkedIn")
    contact.add_run(" | ")
    add_hyperlink(contact, "https://github.com/kutloanom", "GitHub")
    
    # Professional Summary
    doc.add_heading("PROFESSIONAL SUMMARY", level=1)
    summary = doc.add_paragraph()
    run1 = summary.add_run("Full-Stack React.js Developer (BSc Hons Computing, 2025) with proven expertise in ")
    run2 = summary.add_run("React.js, Node.js, TypeScript, JavaScript, and SQL")
    run2.bold = True
    run3 = summary.add_run(". Built production applications serving 100+ users with scalable architectures, RESTful APIs, and optimized database schemas. Creator of AgroSense (PWA), Reusability Compass (analytics platform), and Sesotho AI Platform. UNDP AI Hackathon Finalist with strong foundations in modern development workflows, CI/CD pipelines, and cloud platforms.")
    
    # Technical Skills (Aligned with Job Requirements)
    doc.add_heading("TECHNICAL SKILLS", level=1)
    
    # Frontend Skills
    skill1 = doc.add_paragraph()
    run = skill1.add_run("Frontend Development:")
    run.bold = True
    skill1.add_run(" React.js (18+), TypeScript, JavaScript (ES6+), HTML5, CSS3, Tailwind CSS, Vite, Progressive Web Apps (PWA)")
    
    # Backend Skills
    skill2 = doc.add_paragraph()
    run = skill2.add_run("Backend Development:")
    run.bold = True
    skill2.add_run(" Node.js, Express.js, RESTful APIs, GraphQL (familiar), JWT authentication, API rate limiting, middleware development")
    
    # Database Skills
    skill3 = doc.add_paragraph()
    run = skill3.add_run("Database & SQL:")
    run.bold = True
    skill3.add_run(" PostgreSQL, MySQL, Oracle, database design, indexing, query optimization, Row Level Security (RLS), stored procedures")
    
    # Development Tools
    skill4 = doc.add_paragraph()
    run = skill4.add_run("Development Tools:")
    run.bold = True
    skill4.add_run(" Git, GitHub, CI/CD pipelines, Docker (familiar), automated testing, code reviews, debugging, performance tuning")
    
    # Cloud Platforms
    skill5 = doc.add_paragraph()
    run = skill5.add_run("Cloud Platforms:")
    run.bold = True
    skill5.add_run(" AWS (IoT Core, Lambda, DynamoDB, EC2, S3), Vercel, Render, Supabase, cloud-native architecture")
    
    # Key Projects
    doc.add_heading("KEY PROJECTS", level=1)
    
    # AgroSense Project
    project1 = doc.add_paragraph()
    run = project1.add_run("AgroSense - Smart Agriculture Progressive Web App")
    run.bold = True
    project1.add_run("\\nBotho University (Final Year Project) | ")
    add_hyperlink(project1, "https://agrosense-client-kappa.vercel.app/", "Live Demo")
    project1.add_run(" | ")
    add_hyperlink(project1, "https://github.com/pieter255/AgroSense", "GitHub")
    
    agro_features = [
        "Built scalable React.js frontend with TypeScript, modern hooks, and responsive design serving 100+ farmers",
        "Developed Node.js + Express.js backend with RESTful APIs, JWT authentication, and real-time data processing",
        "Designed PostgreSQL database schema with indexing, optimization, and Row Level Security policies",
        "Implemented offline-first PWA architecture with Service Workers achieving 92% offline availability",
        "Integrated multiple AI APIs with error handling and fallback strategies for high reliability",
        "Established CI/CD pipeline with automated testing, deployment, and monitoring (99.5% uptime)",
        "UNDP AI Hackathon Finalist - competed against established companies, advanced to Phase 2 finals"
    ]
    for feature in agro_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # Reusability Compass Project
    project2 = doc.add_paragraph()
    run = project2.add_run("Reusability Compass - Platform Analytics Dashboard")
    run.bold = True
    project2.add_run("\\nPortfolio Project | ")
    add_hyperlink(project2, "https://reusability-compass.vercel.app", "Live Demo")
    project2.add_run(" | ")
    add_hyperlink(project2, "https://github.com/Shaoza1/reusability-compass", "GitHub")
    
    compass_features = [
        "Built production React.js application with TypeScript, D3.js visualizations, and shadcn/ui components",
        "Developed Node.js + Express.js backend with OpenAI API integration and automated data processing",
        "Implemented Supabase PostgreSQL with real-time synchronization and automated lifecycle management",
        "Created interactive dashboards with complex data visualization and drill-down capabilities",
        "Applied modern development workflows: Git version control, automated testing, performance optimization"
    ]
    for feature in compass_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # Sesotho AI Platform
    project3 = doc.add_paragraph()
    run = project3.add_run("Sesotho AI Platform - Real-time Sentiment Analysis")
    run.bold = True
    project3.add_run("\\nPortfolio Project | ")
    add_hyperlink(project3, "https://github.com/pieter255/moeletsi", "GitHub")
    
    sesotho_features = [
        "Built bilingual React.js frontend with TypeScript, Tailwind CSS, and react-i18next for internationalization",
        "Developed Node.js + Express.js backend with TypeScript, automated web scraping, and real-time data processing",
        "Implemented PostgreSQL database with complex queries, indexing, and secure data storage",
        "Created RESTful APIs with JWT authentication, rate limiting, and comprehensive error handling",
        "Built interactive dashboards with real-time charts and sentiment analysis visualization"
    ]
    for feature in sesotho_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # E-Commerce Platform
    project4 = doc.add_paragraph()
    run = project4.add_run("E-Commerce Web Application")
    run.bold = True
    project4.add_run("\\nAcademic Project (Web Design & Development)")
    
    ecom_features = [
        "Built full-stack e-commerce platform with PHP backend and MySQL database",
        "Implemented user authentication, session management, and role-based access control",
        "Developed shopping cart functionality with secure payment integration and order management",
        "Applied secure coding practices: SQL injection prevention, XSS protection, input validation"
    ]
    for feature in ecom_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # Professional Experience
    doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
    
    # AWS IoT Intern
    exp1 = doc.add_paragraph()
    run = exp1.add_run("Cloud Computing & IoT Virtual Internship")
    run.bold = True
    exp1.add_run("\\nSokul Automation | July - October 2024")
    
    aws_duties = [
        "Developed AWS IoT solutions using Node.js and Python for edge device integration",
        "Built real-time data processing pipelines with cloud-native architecture and monitoring",
        "Applied IAM policies, security best practices, and performance optimization techniques",
        "Collaborated with cross-functional teams on software architecture decisions and code reviews"
    ]
    for duty in aws_duties:
        doc.add_paragraph(f"• {duty}", style='List Bullet')
    
    # Sales Executive
    exp2 = doc.add_paragraph()
    run = exp2.add_run("Sales & Marketing Executive")
    run.bold = True
    exp2.add_run("\\nIMZ Marketing | 2018 - 2019")
    
    sales_duties = [
        "Consistently exceeded sales targets through data-driven analysis and client relationship management",
        "Collaborated with cross-functional teams (marketing, finance, operations) on business solutions",
        "Developed presentations and proposals for executive stakeholders and decision-makers"
    ]
    for duty in sales_duties:
        doc.add_paragraph(f"• {duty}", style='List Bullet')
    
    # Education
    doc.add_heading("EDUCATION", level=1)
    edu = doc.add_paragraph()
    run = edu.add_run("Bachelor of Science (Honours) in Computing")
    run.bold = True
    edu.add_run("\\nBotho University, Lesotho | November 2025 (Graduation: April 2026)")
    
    edu_details = [
        "Core Modules: Data Structures & Algorithms, Software Engineering, Web Design & Development, Database Management",
        "Programming: JavaScript, TypeScript, Java, Python, C++, C#/.NET, PHP, SQL",
        "Specialized: Cloud Computing, Artificial Intelligence, IT Project Management, Network Security",
        "Dean's Award for Academic Excellence (2022, 2024) - Consistent high performance (GPA 3.5+)"
    ]
    for detail in edu_details:
        doc.add_paragraph(f"• {detail}", style='List Bullet')
    
    # Certifications
    doc.add_heading("CERTIFICATIONS", level=1)
    certs = [
        "Python Essentials 1 - Cisco Networking Academy (2025)",
        "Introduction to Cybersecurity - Cisco Networking Academy (2025)",
        "Getting Started with Cisco Packet Tracer - Cisco Networking Academy (2025)"
    ]
    for cert in certs:
        doc.add_paragraph(f"• {cert}", style='List Bullet')
    
    # Awards & Recognition
    doc.add_heading("AWARDS & RECOGNITION", level=1)
    awards = [
        "UNDP AI Language Innovation Hackathon Finalist (2025) - Advanced to Phase 2 with AgroSense platform",
        "Botho University Dean's Award for Academic Achievement (2022, 2024) - Academic excellence recognition"
    ]
    for award in awards:
        doc.add_paragraph(f"• {award}", style='List Bullet')
    
    return doc

def create_iqbusiness_cover_letter():
    """Generate tailored cover letter for IQbusiness React.js Developer position"""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.add_run("Kutloano Peter Moshao\\n")
    header.add_run("Maseru, Lesotho\\n")
    header.add_run("kutloano.moshao111@gmail.com | +266 5758 6176\\n")
    header.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}\\n\\n")
    
    # Recipient
    recipient = doc.add_paragraph()
    recipient.add_run("Hiring Manager\\n")
    recipient.add_run("IQbusiness\\n")
    recipient.add_run("React.js Software Developer Position\\n\\n")
    
    # Salutation
    doc.add_paragraph("Dear Hiring Manager,\\n")
    
    # Opening paragraph
    opening = doc.add_paragraph()
    run1 = opening.add_run("I am writing to express my strong interest in the ")
    run2 = opening.add_run("React.js Software Developer")
    run2.bold = True
    run3 = opening.add_run(" position at IQbusiness. As a recent BSc Honours Computing graduate with proven expertise in React.js, Node.js, TypeScript, JavaScript, and SQL, I have built production applications that demonstrate every technical requirement outlined in your job posting.")
    
    # Technical alignment paragraph
    tech_para = doc.add_paragraph()
    tech_para.add_run("My technical experience directly aligns with your requirements:")
    
    # Technical points
    tech_points = [
        "React.js & TypeScript: Built scalable frontend applications including AgroSense (serving 100+ users) and Reusability Compass with modern React 18, hooks, and TypeScript",
        "Node.js & Express.js: Developed robust backend services with RESTful APIs, JWT authentication, and real-time data processing",
        "SQL Expertise: Designed and optimized PostgreSQL/MySQL schemas with indexing, Row Level Security, and complex query optimization",
        "Modern Workflows: Experienced with Git, CI/CD pipelines, automated testing, code reviews, and performance tuning"
    ]
    for point in tech_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # Project showcase paragraph
    project_para = doc.add_paragraph()
    project_para.add_run("My portfolio showcases production-ready applications: ")
    run1 = project_para.add_run("AgroSense")
    run1.bold = True
    project_para.add_run(" (")
    add_hyperlink(project_para, "https://agrosense-client-kappa.vercel.app/", "live demo")
    project_para.add_run(") - a PWA serving farmers with 92% offline availability and 99.5% uptime, and ")
    run2 = project_para.add_run("Reusability Compass")
    run2.bold = True
    project_para.add_run(" (")
    add_hyperlink(project_para, "https://reusability-compass.vercel.app", "live demo")
    project_para.add_run(") - an analytics platform with D3.js visualizations and real-time data synchronization. Both demonstrate my ability to build scalable, performant, and user-friendly web applications.")
    
    # Experience paragraph
    exp_para = doc.add_paragraph()
    exp_para.add_run("During my AWS IoT internship at Sokul Automation, I developed cloud-native solutions and collaborated with cross-functional teams on software architecture decisions. As a ")
    run = exp_para.add_run("UNDP AI Hackathon Finalist")
    run.bold = True
    exp_para.add_run(", I demonstrated the ability to deliver high-quality software under pressure while maintaining attention to detail and security best practices.")
    
    # Value proposition paragraph
    value_para = doc.add_paragraph()
    value_para.add_run("I am particularly excited about IQbusiness's commitment to sustainable growth and transformation. My experience building applications that handle sensitive data (agricultural records, user authentication, financial transactions) aligns with your emphasis on security and data protection. I bring a passion for clean, reusable code and have consistently applied modern JavaScript/TypeScript best practices across all my projects.")
    
    # Collaboration paragraph
    collab_para = doc.add_paragraph()
    collab_para.add_run("My background includes successful collaboration with diverse teams - from presenting to UNDP officials and international judges to working with technical teams on complex integrations. I understand the importance of cross-functional collaboration with UX/UI, QA, DevOps, and Product teams to deliver exceptional user experiences.")
    
    # Closing paragraph
    closing = doc.add_paragraph()
    closing.add_run("I am eager to contribute to IQbusiness's innovative projects and would welcome the opportunity to discuss how my React.js expertise, full-stack capabilities, and passion for building high-quality software can add value to your development team.")
    
    # Sign-off
    doc.add_paragraph("\\nThank you for your consideration.\\n")
    doc.add_paragraph("Sincerely,\\nKutloano Peter Moshao")
    
    return doc

def main():
    """Generate IQbusiness application documents"""
    print("Generating IQbusiness React.js Developer application documents...")
    
    # Generate CV
    cv_doc = create_iqbusiness_cv()
    cv_filename = "Kutloano_Moshao_CV_IQbusiness_React_Developer.docx"
    cv_doc.save(cv_filename)
    print(f"✅ CV saved as: {cv_filename}")
    
    # Generate Cover Letter
    cl_doc = create_iqbusiness_cover_letter()
    cl_filename = "Kutloano_Moshao_Cover_Letter_IQbusiness_React.docx"
    cl_doc.save(cl_filename)
    print(f"✅ Cover Letter saved as: {cl_filename}")
    
    print("\\n🎯 Documents tailored for IQbusiness React.js Software Developer")
    print("📋 Key highlights:")
    print("   • React.js + TypeScript + Node.js expertise demonstrated")
    print("   • Production applications with live demos")
    print("   • SQL database design and optimization experience")
    print("   • Modern development workflows and CI/CD")
    print("   • Cross-functional collaboration experience")

if __name__ == "__main__":
    main()