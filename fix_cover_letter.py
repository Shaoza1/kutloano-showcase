#!/usr/bin/env python3
"""
Generate FIXED cover letter for IQbusiness React.js Developer position
Properly handles line breaks in Word documents without showing \\n
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import datetime

def add_hyperlink(paragraph, url, text):
    """Add hyperlink to paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0563C1")
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_fixed_cover_letter():
    """Generate properly formatted cover letter without \\n characters"""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.add_run("Kutloano Peter Moshao")
    header.add_run().add_break()
    header.add_run("Maseru, Lesotho")
    header.add_run().add_break()
    header.add_run("kutloano.moshao111@gmail.com | +266 5758 6176")
    header.add_run().add_break()
    header.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
    
    # Add spacing
    doc.add_paragraph()
    
    # Recipient
    recipient = doc.add_paragraph()
    recipient.add_run("Hiring Manager")
    recipient.add_run().add_break()
    recipient.add_run("IQbusiness")
    recipient.add_run().add_break()
    recipient.add_run("React.js Software Developer Position")
    
    # Add spacing
    doc.add_paragraph()
    
    # Salutation
    doc.add_paragraph("Dear Hiring Manager,")
    
    # Opening paragraph
    opening = doc.add_paragraph()
    run1 = opening.add_run("I am writing to express my strong interest in the ")
    run2 = opening.add_run("React.js Software Developer")
    run2.bold = True
    run3 = opening.add_run(" position at IQbusiness. As a recent BSc Honours Computing graduate with proven expertise in React.js, Node.js, TypeScript, JavaScript, and SQL, I have built production applications that demonstrate every technical requirement outlined in your job posting.")
    
    # Technical alignment paragraph
    tech_para = doc.add_paragraph()
    tech_para.add_run("My technical experience directly aligns with your requirements:")
    
    # Technical points
    tech_points = [
        "React.js & TypeScript: Built scalable frontend applications including AgroSense (serving 100+ users) and Reusability Compass with modern React 18, hooks, and TypeScript",
        "Node.js & Express.js: Developed robust backend services with RESTful APIs, JWT authentication, and real-time data processing",
        "SQL Expertise: Designed and optimized PostgreSQL/MySQL schemas with indexing, Row Level Security, and complex query optimization",
        "Modern Workflows: Experienced with Git, CI/CD pipelines, automated testing, code reviews, and performance tuning"
    ]
    for point in tech_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # Project showcase paragraph
    project_para = doc.add_paragraph()
    project_para.add_run("My portfolio showcases production-ready applications: ")
    run1 = project_para.add_run("AgroSense")
    run1.bold = True
    project_para.add_run(" (")
    add_hyperlink(project_para, "https://agrosense-client-kappa.vercel.app/", "live demo")
    project_para.add_run(") - a PWA serving farmers with 92% offline availability and 99.5% uptime, and ")
    run2 = project_para.add_run("Reusability Compass")
    run2.bold = True
    project_para.add_run(" (")
    add_hyperlink(project_para, "https://reusability-compass.vercel.app", "live demo")
    project_para.add_run(") - an analytics platform with D3.js visualizations and real-time data synchronization. Both demonstrate my ability to build scalable, performant, and user-friendly web applications.")
    
    # Experience paragraph
    exp_para = doc.add_paragraph()
    exp_para.add_run("During my AWS IoT internship at Sokul Automation, I developed cloud-native solutions and collaborated with cross-functional teams on software architecture decisions. As a ")
    run = exp_para.add_run("UNDP AI Hackathon Finalist")
    run.bold = True
    exp_para.add_run(", I demonstrated the ability to deliver high-quality software under pressure while maintaining attention to detail and security best practices.")
    
    # Value proposition paragraph
    value_para = doc.add_paragraph()
    value_para.add_run("I am particularly excited about IQbusiness's commitment to sustainable growth and transformation. My experience building applications that handle sensitive data (agricultural records, user authentication, financial transactions) aligns with your emphasis on security and data protection. I bring a passion for clean, reusable code and have consistently applied modern JavaScript/TypeScript best practices across all my projects.")
    
    # Collaboration paragraph
    collab_para = doc.add_paragraph()
    collab_para.add_run("My background includes successful collaboration with diverse teams - from presenting to UNDP officials and international judges to working with technical teams on complex integrations. I understand the importance of cross-functional collaboration with UX/UI, QA, DevOps, and Product teams to deliver exceptional user experiences.")
    
    # Closing paragraph
    closing = doc.add_paragraph()
    closing.add_run("I am eager to contribute to IQbusiness's innovative projects and would welcome the opportunity to discuss how my React.js expertise, full-stack capabilities, and passion for building high-quality software can add value to your development team.")
    
    # Add spacing
    doc.add_paragraph()
    
    # Thank you
    doc.add_paragraph("Thank you for your consideration.")
    
    # Add spacing
    doc.add_paragraph()
    
    # Sign-off
    signoff = doc.add_paragraph()
    signoff.add_run("Sincerely,")
    signoff.add_run().add_break()
    signoff.add_run("Kutloano Peter Moshao")
    
    return doc

def main():
    """Generate fixed cover letter"""
    print("Generating FIXED IQbusiness cover letter...")
    
    # Generate Fixed Cover Letter
    cl_doc = create_fixed_cover_letter()
    cl_filename = "Kutloano_Moshao_FIXED_Cover_Letter_IQbusiness.docx"
    cl_doc.save(cl_filename)
    print(f"✅ Fixed Cover Letter saved as: {cl_filename}")
    
    print("\\n🎯 Fixed cover letter - no more \\\\n characters in Word document!")

if __name__ == "__main__":
    main()