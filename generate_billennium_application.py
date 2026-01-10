#!/usr/bin/env python3
"""
Generate comprehensive tailored CV and Cover Letter for Billennium Software Engineer position
Includes ALL projects and experiences with strategic positioning of Reusability Compass
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import datetime

def add_hyperlink(paragraph, url, text):
    """Add hyperlink to paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0563C1")
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_comprehensive_cv():
    """Generate comprehensive CV with all projects and strategic Billennium positioning"""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("KUTLOANO PETER MOSHAO")
    run.font.size = Pt(16)
    run.bold = True
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("Maseru, Lesotho | kutloano.moshao111@gmail.com | +266 5758 6176\n")
    add_hyperlink(contact, "https://kutloano-showcase.vercel.app", "kutloano-showcase.vercel.app")
    contact.add_run(" | ")
    add_hyperlink(contact, "https://linkedin.com/in/kutloano-moshao-1aa5003a1", "LinkedIn")
    contact.add_run(" | ")
    add_hyperlink(contact, "https://github.com/kutloanom", "GitHub")
    
    # Professional Summary
    doc.add_heading("SUMMARY", level=1)
    summary = doc.add_paragraph()
    summary.add_run("I completed my Bachelor of Science (Honours) in Computing at Botho University in November 2025 and am awaiting graduation in April 2026. I am a Lesotho national currently without a South African work permit. I am prepared to apply for a General Work Visa / Critical Skills Visa immediately upon receipt of a formal job offer. ")
    
    run1 = summary.add_run("Software Engineer (BSc Hons Computing, 2025) specializing in survey automation, GenAI evaluation, and interactive data visualization")
    run1.bold = True
    summary.add_run(". Creator of ")
    run2 = summary.add_run("Reusability Compass")
    run2.bold = True
    summary.add_run(" - production platform demonstrating exact Billennium role requirements. Experienced in building secure, scalable applications with AI integration. Delivered AgroSense (PWA), Sesotho AI Platform, and enterprise analytics solutions. UNDP AI Hackathon Finalist with strong foundations in full-stack development, machine learning, and enterprise software architecture.")
    
    # Experience
    doc.add_heading("EXPERIENCE", level=1)
    
    # AWS IoT Intern
    exp1 = doc.add_paragraph()
    run = exp1.add_run("Cloud Computing & IoT Virtual Internship")
    run.bold = True
    exp1.add_run("\nSokul Automation")
    exp1.add_run("\nJuly 2024 - October 2024, Redwood City, USA")
    
    aws_duties = [
        "Developed AWS IoT Greengrass components in Python for edge device integration",
        "Built real-time data processing pipelines and automation scripts", 
        "Applied IAM policies for secure access control and least-privilege principles",
        "Provisioned EC2 instances and implemented cloud-native monitoring solutions",
        "Recognized for persistence, leadership, and technical aptitude"
    ]
    for duty in aws_duties:
        doc.add_paragraph(f"• {duty}", style='List Bullet')
    
    # Projects Section
    doc.add_heading("PROJECTS", level=1)
    
    # Reusability Compass (Featured for Billennium)
    project1 = doc.add_paragraph()
    run = project1.add_run("Reusability Compass – Platform Insights & Adoption Friction Analyzer")
    run.bold = True
    project1.add_run("\nPortfolio Project • ")
    add_hyperlink(project1, "https://reusability-compass.vercel.app", "reusability-compass.vercel.app")
    project1.add_run(" | ")
    add_hyperlink(project1, "https://github.com/Shaoza1/reusability-compass", "GitHub")
    
    compass_features = [
        "Built production-grade platform demonstrating EXACT Billennium role requirements: survey automation, GenAI evaluation, interactive visualization",
        "Survey Automation: Automated developer feedback collection with structured data capture and real-time validation",
        "GenAI Evaluation: OpenAI GPT-4 integration for sentiment analysis, friction scoring (0-1 scale), and root cause extraction", 
        "Interactive Visualization: D3.js heatmap displaying many-to-many platform-project relationships with drill-down capabilities",
        "Data Lifecycle Management: Automated survey → AI analysis → visualization pipeline with status tracking and triggers",
        "Architecture: React 18 + TypeScript frontend, Node.js + Express backend, Supabase PostgreSQL with Row Level Security",
        "Enterprise Features: Real-time data synchronization, comprehensive error handling, executive dashboard for leadership",
        "Business Impact: Quantifies platform adoption friction and strategic value for engineering teams"
    ]
    for feature in compass_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # AgroSense
    project2 = doc.add_paragraph()
    run = project2.add_run("AgroSense – Smart Agriculture Progressive Web App")
    run.bold = True
    project2.add_run("\nBotho University (Final Year Project) • ")
    add_hyperlink(project2, "https://agrosense-client-kappa.vercel.app/", "agrosense-client-kappa.vercel.app")
    
    agro_features = [
        "UNDP AI Hackathon Finalist - Advanced to Phase 2 finals competing against established companies",
        "Integrated multiple AI APIs (OpenAI GPT-4, Google Gemini, Cohere) for crop disease detection with 89.3% accuracy",
        "Designed PostgreSQL/MySQL schemas with indexing, optimization, and row-level security policies",
        "Implemented offline-first architecture with Service Workers and IndexedDB achieving 92% offline availability",
        "Built Sesotho voice command system with fuzzy keyword matching for rural farmers",
        "Established CI/CD pipeline, automated backups, monitoring, and audit logging",
        "Deployed production system serving 100+ farmers with 94% adoption rate and 76.8 SUS usability score",
        "Tech Stack: React 18, TypeScript, Node.js, Supabase, Capacitor, PWA, Multi-AI integration"
    ]
    for feature in agro_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # Sesotho AI Platform
    project3 = doc.add_paragraph()
    run = project3.add_run("Sesotho AI Platform (Moeletsi)")
    run.bold = True
    project3.add_run("\nPortfolio Project • ")
    add_hyperlink(project3, "https://github.com/pieter255/moeletsi", "GitHub")
    
    sesotho_features = [
        "Designed real-time sentiment analysis for Lesotho news articles using gold-standard Sesotho News Dataset",
        "Built bilingual frontend (React + TypeScript) with Tailwind CSS and react-i18next for Sesotho/English support",
        "Developed backend services (Node.js + Express + TypeScript) with Supabase PostgreSQL for secure data storage",
        "Created ML service (Python Flask + scikit-learn) for sentiment scoring, topic extraction, and aspect analysis",
        "Integrated AI models (OpenAI GPT, Google Gemini, Groq LLaMA) for Retrieval-Augmented Generation (RAG) assistant",
        "Automated web scraping with Python (BeautifulSoup, Selenium) for Lesotho news sources",
        "Implemented security: JWT authentication, row-level security (RLS), API rate limiting, input validation",
        "Built interactive dashboards for sentiment trends, distribution charts, and topic insights with monitoring"
    ]
    for feature in sesotho_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # E-Commerce Platform
    project4 = doc.add_paragraph()
    run = project4.add_run("E-Commerce Web Application")
    run.bold = True
    project4.add_run("\nAcademic Project (Web Design & Development Module)")
    
    ecom_features = [
        "Built full-stack e-commerce platform with PHP and MySQL demonstrating enterprise web development",
        "Implemented user authentication, session management, and role-based access control",
        "Delivered shopping cart functionality, payment integration, and comprehensive order management",
        "Applied secure coding practices: SQL injection prevention, XSS protection, password hashing",
        "Designed normalized database schema with foreign keys, indexes, and stored procedures"
    ]
    for feature in ecom_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # Education
    doc.add_heading("EDUCATION", level=1)
    edu = doc.add_paragraph()
    run = edu.add_run("Bachelor of Science (Honours) in Computing")
    run.bold = True
    edu.add_run("\nBotho University, Lesotho • November 2025 (Graduation Ceremony: April 2026)")
    
    edu_details = [
        "Core Modules: Data Structures & Algorithms, Operating Systems, Software Engineering, Cloud Computing, Artificial Intelligence, Machine Learning",
        "Programming Languages: Python, Java, C++, C#/.NET, PHP, JavaScript, TypeScript",
        "Specialized Coursework: Web Design & Development, Database Management, Network Security, Mobile Application Development",
        "22 Total Modules: Introduction to Computer, Mathematics for Computing, Communication Skills, Database Management (Oracle), Linux Essentials, IT Project Management, Advanced Java, Programming using .Net, Interaction Design, Essentials of Entrepreneurship",
        "Awarded Dean's Award for Academic Achievement three times (2022, 2024) - GPA 3.5+"
    ]
    for detail in edu_details:
        doc.add_paragraph(f"• {detail}", style='List Bullet')
    
    # Certifications
    doc.add_heading("CERTIFICATIONS", level=1)
    
    cert1 = doc.add_paragraph()
    run = cert1.add_run("Python Essentials 1")
    run.bold = True
    cert1.add_run("\nCisco Networking Academy • 2025")
    cert1_details = [
        "Strengthened Python programming skills, backend scripting, and automation capabilities",
        "Directly relevant to Billennium's requirement for automation scripting and data processing pipelines",
        "Applied in AWS IoT Greengrass development and backend automation projects"
    ]
    for detail in cert1_details:
        doc.add_paragraph(f"• {detail}", style='List Bullet')
    
    cert2 = doc.add_paragraph()
    run = cert2.add_run("Introduction to Cybersecurity")
    run.bold = True
    cert2.add_run("\nCisco Networking Academy • 2025")
    cert2_details = [
        "Gained knowledge of secure coding practices, encryption, authentication, and data protection principles",
        "Relevant to Billennium's enterprise security requirements for survey data and GenAI pipeline protection",
        "Applied in implementing JWT authentication, RLS policies, and secure API development"
    ]
    for detail in cert2_details:
        doc.add_paragraph(f"• {detail}", style='List Bullet')
    
    cert3 = doc.add_paragraph()
    run = cert3.add_run("Getting Started with Cisco Packet Tracer")
    run.bold = True
    cert3.add_run("\nCisco Networking Academy • 2025")
    cert3_details = [
        "Developed practical networking and troubleshooting skills for distributed systems",
        "Relevant for diagnosing connectivity issues in complex data visualization and real-time systems",
        "Understanding of multi-tiered systems architecture essential for enterprise platform analytics"
    ]
    for detail in cert3_details:
        doc.add_paragraph(f"• {detail}", style='List Bullet')
    
    # Awards & Honors
    doc.add_heading("AWARDS & HONORS", level=1)
    
    award1 = doc.add_paragraph()
    run = award1.add_run("UNDP AI Language Innovation Hackathon – Finalist")
    run.bold = True
    award1.add_run("\nUNDP Lesotho • December 2025")
    award1_details = [
        "Competed as solo entrant against established companies, advancing to Phase 2 finals",
        "Presented AgroSense platform to international judges and UNDP officials",
        "Demonstrated ability to innovate with AI under pressure - exactly what Billennium values",
        "Showcased skills in articulating technical challenges and solutions to executive stakeholders",
        "Proved capability to apply GenAI concepts to real-world problems with scalable architecture"
    ]
    for detail in award1_details:
        doc.add_paragraph(f"• {detail}", style='List Bullet')
    
    award2 = doc.add_paragraph()
    run = award2.add_run("Botho University Dean's Award for Academic Achievement")
    run.bold = True
    award2.add_run("\nBotho University • 2022, 2024")
    award2_details = [
        "Maintained GPA of 3.5+ across multiple semesters while balancing technical projects and innovation",
        "Demonstrates consistent academic excellence, discipline, and mastery of computer science fundamentals",
        "Shows operational discipline and persistence that Billennium expects in software engineers",
        "Reinforces ability to learn quickly, adapt, and sustain high performance in fast-paced environments"
    ]
    for detail in award2_details:
        doc.add_paragraph(f"• {detail}", style='List Bullet')
    
    # Technical Skills (Organized by Billennium Requirements)
    doc.add_heading("TECHNICAL SKILLS", level=1)
    
    # Survey Automation & Data Collection
    skill1 = doc.add_paragraph()
    run = skill1.add_run("Survey Automation & Data Collection:")
    run.bold = True
    skill1.add_run(" React Hook Form, Zod validation, automated workflows, survey tool integration, structured data capture")
    
    # GenAI & AI Integration
    skill2 = doc.add_paragraph()
    run = skill2.add_run("GenAI & AI Integration:")
    run.bold = True
    skill2.add_run(" OpenAI GPT-4, Google Gemini, Cohere, Groq, HuggingFace, Prompt Engineering, NLP, Sentiment Analysis, TensorFlow (familiar), PyTorch (familiar)")
    
    # Data Visualization & UX
    skill3 = doc.add_paragraph()
    run = skill3.add_run("Data Visualization & UX:")
    run.bold = True
    skill3.add_run(" D3.js, Recharts, React 18, TypeScript, Tailwind CSS, shadcn/ui, interactive dashboards, heatmap visualization, many-to-many relationship representation")
    
    # Backend & Database
    skill4 = doc.add_paragraph()
    run = skill4.add_run("Backend & Database:")
    run.bold = True
    skill4.add_run(" Node.js, Express.js, Python Flask, PostgreSQL, MySQL, Oracle, Supabase, database design, indexing, optimization, Row Level Security (RLS)")
    
    # Programming Languages
    skill5 = doc.add_paragraph()
    run = skill5.add_run("Programming Languages:")
    run.bold = True
    skill5.add_run(" Python, JavaScript, TypeScript, Java, C++, C#/.NET, PHP, SQL, Bash")
    
    # Cloud & DevOps
    skill6 = doc.add_paragraph()
    run = skill6.add_run("Cloud & DevOps:")
    run.bold = True
    skill6.add_run(" AWS (IoT Core, Lambda, DynamoDB, EC2, S3), Vercel, Render, Docker (familiar), CI/CD pipelines, automated backups, monitoring")
    
    # Security & Authentication
    skill7 = doc.add_paragraph()
    run = skill7.add_run("Security & Authentication:")
    run.bold = True
    skill7.add_run(" JWT authentication, SQL injection prevention, XSS protection, CSRF awareness, password hashing, API rate limiting, audit logging")
    
    return doc

def create_comprehensive_cover_letter():
    """Generate comprehensive cover letter highlighting all relevant experience"""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.add_run("Kutloano Peter Moshao\n")
    header.add_run("Maseru, Lesotho\n")
    header.add_run("kutloano.moshao111@gmail.com | +266 5758 6176\n")
    header.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}\n\n")
    
    # Recipient
    recipient = doc.add_paragraph()
    recipient.add_run("Hiring Manager\n")
    recipient.add_run("Billennium\n")
    recipient.add_run("Software Engineer - Reusability Insights & Visualization\n\n")
    
    # Salutation
    doc.add_paragraph("Dear Hiring Manager,\n")
    
    # Opening - Strong hook with Reusability Compass
    opening = doc.add_paragraph()
    run1 = opening.add_run("I am writing to express my strong interest in the ")
    run2 = opening.add_run("Software Engineer - Reusability Insights & Visualization")
    run2.bold = True
    run3 = opening.add_run(" position at Billennium. I have built ")
    run4 = opening.add_run("Reusability Compass")
    run4.bold = True
    run5 = opening.add_run(" (")
    add_hyperlink(opening, "https://reusability-compass.vercel.app", "live demo")
    run6 = opening.add_run(") - a production platform that demonstrates every single requirement in your job posting: survey automation, GenAI evaluation, interactive visualization, and data lifecycle management.")
    
    # Perfect alignment paragraph
    alignment = doc.add_paragraph()
    alignment.add_run("Your role requirements align perfectly with my demonstrated expertise:")
    
    # Alignment bullets
    align_points = [
        "Survey Automation: Built automated developer feedback collection with structured data capture and real-time validation in Reusability Compass",
        "GenAI Evaluation: Implemented OpenAI GPT-4 pipelines for sentiment analysis, friction scoring, and root cause extraction from unstructured survey data",
        "Interactive Visualization: Created D3.js heatmaps displaying complex many-to-many platform relationships with drill-down capabilities and UX design for leadership consumption",
        "Data Lifecycle Management: Developed automated survey → AI analysis → visualization pipeline with status tracking and database triggers"
    ]
    for point in align_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # Comprehensive experience paragraph
    experience = doc.add_paragraph()
    experience.add_run("Beyond Reusability Compass, my portfolio demonstrates comprehensive software engineering capabilities. As a ")
    run = experience.add_run("UNDP AI Hackathon Finalist")
    run.bold = True
    experience.add_run(", I competed solo against established companies with AgroSense - an AI-powered agricultural platform serving 100+ farmers with 94% adoption rate. During my AWS IoT internship at Sokul Automation, I developed cloud-native solutions with Python, implemented real-time data pipelines, and applied enterprise security practices. My Sesotho AI Platform showcases advanced NLP capabilities, automated web scraping, and bilingual interface design.")
    
    # Technical depth paragraph
    technical = doc.add_paragraph()
    technical.add_run("My technical foundation spans the full stack: React 18 + TypeScript + D3.js for interactive frontends, Node.js + Express for scalable backends, PostgreSQL with Row Level Security for enterprise data management, and comprehensive AI integration (OpenAI GPT-4, Google Gemini, HuggingFace). I have production experience with automated workflows, CI/CD pipelines, real-time synchronization, and enterprise-grade error handling - all essential for Billennium's platform efficiency initiatives.")
    
    # Academic excellence paragraph
    academic = doc.add_paragraph()
    academic.add_run("My BSc Honours in Computing (2025) from Botho University, where I earned the Dean's Award three times, provided strong foundations in data structures, algorithms, software engineering, and AI. My Cisco certifications in Python programming and cybersecurity demonstrate commitment to continuous learning and security-first development practices.")
    
    # Value proposition paragraph
    value = doc.add_paragraph()
    value.add_run("I am particularly excited about Billennium's 20+ years of innovation and global collaboration culture. My experience transforming qualitative developer feedback into structured insights, building executive dashboards for strategic decision-making, and creating comprehensive documentation for team handoffs aligns perfectly with your mission to empower businesses through technology.")
    
    # Closing paragraph
    closing = doc.add_paragraph()
    closing.add_run("I have demonstrated the exact capabilities you seek through production deployments, not just theoretical knowledge. I am eager to contribute to Billennium's platform efficiency initiatives and would welcome the opportunity to discuss how my specialized experience in survey automation, GenAI evaluation, and interactive visualization can drive strategic value for your engineering teams.")
    
    # Sign-off
    doc.add_paragraph("\nThank you for your consideration.\n")
    doc.add_paragraph("Sincerely,\nKutloano Peter Moshao")
    
    return doc

def main():
    """Generate comprehensive CV and cover letter"""
    print("Generating comprehensive Billennium application documents...")
    
    # Generate Comprehensive CV
    cv_doc = create_comprehensive_cv()
    cv_filename = "Kutloano_Moshao_COMPREHENSIVE_CV_Billennium.docx"
    cv_doc.save(cv_filename)
    print(f"✅ Comprehensive CV saved as: {cv_filename}")
    
    # Generate Comprehensive Cover Letter
    cl_doc = create_comprehensive_cover_letter()
    cl_filename = "Kutloano_Moshao_COMPREHENSIVE_Cover_Letter_Billennium.docx"
    cl_doc.save(cl_filename)
    print(f"✅ Comprehensive Cover Letter saved as: {cl_filename}")
    
    print("\n🎯 COMPREHENSIVE Documents for Billennium Software Engineer Position")
    print("📋 Includes ALL projects and experiences:")
    print("   • Reusability Compass (perfect role match)")
    print("   • AgroSense (UNDP Hackathon Finalist)")
    print("   • Sesotho AI Platform (advanced NLP)")
    print("   • E-Commerce Platform (full-stack web dev)")
    print("   • AWS IoT experience (cloud-native solutions)")
    print("   • All certifications and academic achievements")
    print("   • Strategic positioning for maximum impact")

if __name__ == "__main__":
    main()